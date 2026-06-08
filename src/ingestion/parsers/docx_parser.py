"""
Word (.docx) 文档解析器。

使用 python-docx 提取段落文本，保留章节标题层级信息。
每个段落作为一个 LlamaDocument。
"""

import logging
from typing import List

from llama_index.core.schema import Document as LlamaDocument

from src.ingestion.parsers.base_parser import BaseParser, make_doc_id

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Word 文档解析器——基于 python-docx。

    提取所有段落文本，识别标题样式（Heading 1-3），
    记录每个段落的当前章节信息。

    Usage::

        parser = DocxParser()
        docs = parser.parse("/path/to/doc.docx")
    """

    @property
    def supported_format(self) -> str:
        return "docx"

    def parse(self, source: str) -> List[LlamaDocument]:
        """解析 .docx 文件。

        Args:
            source: .docx 文件路径。

        Returns:
            LlamaDocument 列表，每段一个文档。
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx 未安装")
            return []

        doc_id = make_doc_id(source)
        doc_name = self._source_name(source)
        documents: List[LlamaDocument] = []

        try:
            docx = DocxDocument(source)
        except Exception as e:
            logger.error("无法打开 Word 文档 %s: %s", source, e)
            return []

        # 跟踪当前章节上下文
        current_heading = ""
        paragraph_index = 0
        total_chars = 0

        for para in docx.paragraphs:
            text = para.text
            if not text or not text.strip():
                # 更新标题上下文（空标题段落也需跟踪）
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    current_heading = text.strip() if text.strip() else current_heading
                continue

            total_chars += len(text)

            # 检测标题样式
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                current_heading = text.strip()
                # 标题本身也作为一个 chunk（简短但包含重要语义）
                doc = LlamaDocument(
                    text=text.strip(),
                    metadata={
                        "doc_id": doc_id,
                        "doc_name": doc_name,
                        "source": source,
                        "format": "docx",
                        "heading": current_heading,
                        "paragraph_index": paragraph_index,
                        "is_heading": True,
                    },
                )
                documents.append(doc)
                paragraph_index += 1
                continue

            doc = LlamaDocument(
                text=text.strip(),
                metadata={
                    "doc_id": doc_id,
                    "doc_name": doc_name,
                    "source": source,
                    "format": "docx",
                    "heading": current_heading or "",
                    "paragraph_index": paragraph_index,
                    "is_heading": False,
                },
            )
            documents.append(doc)
            paragraph_index += 1

        logger.info(
            "Word 解析完成: %s (%d 段, %d 字符)", doc_name, len(documents), total_chars
        )
        return documents

    # ------------------------------------------------------------------
    # LlamaIndex 回退
    # ------------------------------------------------------------------

    @staticmethod
    def parse_via_llamaindex(source: str) -> List[LlamaDocument]:
        """使用 LlamaIndex DocxReader 作为备选方案。

        Args:
            source: 文件路径。

        Returns:
            LlamaDocument 列表。
        """
        try:
            from llama_index.readers.file import DocxReader

            reader = DocxReader()
            docs = reader.load_data(file_path=source)  # type: ignore[arg-type]
            logger.info("LlamaIndex DocxReader 解析完成: %s", source)
            return list(docs)
        except ImportError:
            logger.error("llama-index-readers-file 未安装")
            return []
